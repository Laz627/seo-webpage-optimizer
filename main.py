import streamlit as st
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from io import BytesIO

# Set page config
st.set_page_config(page_title="Content Optimizer", layout="wide")
st.title("Content Optimizer")

# Instructions
st.markdown("""
## How to Use This Tool:
1. **Enter your OpenAI API key** in the field below.
2. **Upload the HTML file** of your web page.
3. **Input your target keyword.**
4. **Upload competitor HTML files** for comparison.
5. Click **'Optimize Content'** to analyze competitor pages and receive content recommendations.
6. **Download the recommendations** as a Word document.

This tool provides specific recommendations on how to enhance your content based on competitor analysis, including new sections to add, where to place them, the appropriate heading levels, and suggestions for your meta title, meta description, and H1 tag based on SEO best practices.

**Note:** The script ignores irrelevant content such as menu navigation links, footer links, and anchor text not part of the main content.
""")

# Initialize session state
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = ''
if 'keyword' not in st.session_state:
    st.session_state.keyword = ''

def extract_content_structure(html_content):
    try:
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove navigation and footer elements
        tags_to_remove = ['script', 'style', 'noscript', 'header', 'footer', 'nav', 'aside']
        for tag in tags_to_remove:
            for element in soup.find_all(tag):
                element.decompose()

        # Remove elements by common class and ID names related to navigation, sidebar, and footer
        classes_ids_to_remove = ['nav', 'navigation', 'sidebar', 'footer', 'header', 'menu', 'breadcrumbs', 'breadcrumb', 'site-footer', 'site-header', 'widget', 'widgets', 'site-navigation', 'main-navigation', 'secondary-navigation', 'site-sidebar']
        for class_or_id in classes_ids_to_remove:
            for element in soup.find_all(attrs={'class': class_or_id}):
                element.decompose()
            for element in soup.find_all(attrs={'id': class_or_id}):
                element.decompose()

        # Extract meta title and description
        meta_title = soup.title.string.strip() if soup.title else ''
        meta_description_tag = soup.find('meta', attrs={'name': 'description'})
        meta_description = meta_description_tag['content'].strip() if meta_description_tag else ''
        # Extract H1 tag
        h1_tag = soup.find('h1')
        h1_text = h1_tag.get_text(separator=' ', strip=True) if h1_tag else ''

        # Remove any remaining scripts and styles
        for element in soup(['script', 'style', 'noscript']):
            element.decompose()

        # Extract headings and their hierarchy
        content_structure = []
        for header in soup.find_all(['h2', 'h3', 'h4']):
            # Exclude headings within navigation elements
            if header.find_parent(['nav', 'header', 'footer', 'aside']):
                continue
            if 'class' in header.attrs and any(cls in classes_ids_to_remove for cls in header.get('class', [])):
                continue
            if 'id' in header.attrs and header['id'] in classes_ids_to_remove:
                continue

            header_text = header.get_text(separator=' ', strip=True)
            header_level = header.name.upper()
            content_structure.append({'level': header_level, 'text': header_text})
        return meta_title, meta_description, h1_text, content_structure
    except Exception as e:
        st.warning(f"Error extracting content structure: {str(e)}")
        return '', '', '', []

def analyze_competitor_content(html_files):
    all_headings = []
    for file in html_files:
        try:
            html_content = file.read().decode('utf-8')
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove navigation and footer elements
            tags_to_remove = ['script', 'style', 'noscript', 'header', 'footer', 'nav', 'aside']
            for tag in tags_to_remove:
                for element in soup.find_all(tag):
                    element.decompose()

            # Remove elements by common class and ID names
            classes_ids_to_remove = ['nav', 'navigation', 'sidebar', 'footer', 'header', 'menu', 'breadcrumbs', 'breadcrumb', 'site-footer', 'site-header', 'widget', 'widgets', 'site-navigation', 'main-navigation', 'secondary-navigation', 'site-sidebar']
            for class_or_id in classes_ids_to_remove:
                for element in soup.find_all(attrs={'class': class_or_id}):
                    element.decompose()
                for element in soup.find_all(attrs={'id': class_or_id}):
                    element.decompose()

            # Extract meta title and description
            meta_title = soup.title.string.strip() if soup.title else ''
            meta_description_tag = soup.find('meta', attrs={'name': 'description'})
            meta_description = meta_description_tag['content'].strip() if meta_description_tag else ''
            # Extract H1 tag
            h1_tag = soup.find('h1')
            h1_text = h1_tag.get_text(separator=' ', strip=True) if h1_tag else ''

            # Remove any remaining scripts and styles
            for element in soup(['script', 'style', 'noscript']):
                element.decompose()

            # Extract headings
            headings = []
            for header in soup.find_all(['h2', 'h3', 'h4']):
                # Exclude headings within navigation elements
                if header.find_parent(['nav', 'header', 'footer', 'aside']):
                    continue
                if 'class' in header.attrs and any(cls in classes_ids_to_remove for cls in header.get('class', [])):
                    continue
                if 'id' in header.attrs and header['id'] in classes_ids_to_remove:
                    continue

                header_text = header.get_text(separator=' ', strip=True)
                header_level = header.name.upper()
                headings.append({'level': header_level, 'text': header_text})
            competitor_data = {
                'meta_title': meta_title,
                'meta_description': meta_description,
                'h1_text': h1_text,
                'headings': headings
            }
            all_headings.append(competitor_data)
        except Exception as e:
            st.warning(f"Error processing {file.name}: {e}")
            continue
    return all_headings

def generate_detailed_recommendations(keyword, meta_title, meta_description, h1_text, user_structure, competitor_data, api_key):
    client = OpenAI(api_key=api_key)

    # Prepare competitor meta data
    competitor_meta_info = ''
    for idx, comp in enumerate(competitor_data, 1):
        competitor_meta_info += f"Competitor #{idx} Meta Title: {comp['meta_title']}\n"
        competitor_meta_info += f"Competitor #{idx} Meta Description: {comp['meta_description']}\n"
        competitor_meta_info += f"Competitor #{idx} H1 Tag: {comp['h1_text']}\n"
        competitor_headings_str = '\n'.join([f"{item['level']}: {item['text']}" for item in comp['headings']])
        competitor_meta_info += f"Competitor #{idx} Headings:\n{competitor_headings_str}\n\n"

    prompt = f"""
You are an SEO content strategist.

Your task is to analyze the provided original content structure, meta information, and competitor data to generate specific recommendations for improvement.

- **Keyword**: "{keyword}"
- **Original Meta Title**: {meta_title}
- **Original Meta Description**: {meta_description}
- **Original H1 Tag**: {h1_text}
- **Original Content Structure**:
{user_structure}

- **Competitor Meta Information and Headings**:
{competitor_meta_info}

Instructions:

1. Analyze the original meta title, meta description, and H1 tag. Provide specific recommendations for improving each based on SEO best practices and the target keyword.
2. Identify important topics or subtopics in the competitor headings that are missing or underdeveloped in the original content.
3. For each identified area:
   - Recommend new sections to add.
   - Specify where to place them within the existing content structure.
   - Indicate the appropriate heading level (H2/H3/H4).
   - Provide a suggested heading title.
   - Briefly describe what content should be included under each new section.
4. If rearranging existing sections would improve content flow, provide specific suggestions.
5. If the original content is already comprehensive, acknowledge that but suggest any minor improvements if applicable.
6. Avoid using branded terms unless necessary.
7. Present the recommendations clearly and in a structured format.

Format:

Meta Title Recommendation:
[Your recommendation here]

Meta Description Recommendation:
[Your recommendation here]

H1 Tag Recommendation:
[Your recommendation here]

Content Recommendations:

For each recommendation:

Recommendation #X:
- New Section Heading: [Suggested heading title]
- Heading Level: [H2/H3/H4]
- Placement: [Where to insert in the existing structure]
- Content Description: [Brief description of what to include]

If rearranging sections:

Rearrangement Suggestion:
- Current Section: [Existing heading]
- New Position: [Where it should be moved]
- Reason: [Why this improves the content flow]

Provide a final summary acknowledging if the content is comprehensive or noting any overall improvements.

IMPORTANT: Do not include any additional text outside of the specified format. Do not use bullet points or leading hyphens in your output.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Provide detailed SEO content recommendations based on the analysis."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        output = response.choices[0].message.content
        return output
    except Exception as e:
        st.error(f"Error generating recommendations:\n\n{str(e)}")
        return None

def parse_recommendations(recommendations_text):
    # Parse the recommendations into structured data
    recommendations = {'meta_title': '', 'meta_description': '', 'h1_tag': '', 'content': [], 'rearrangements': [], 'summary': ''}
    lines = recommendations_text.strip().split('\n')
    section = None
    content_rec = {}
    for line in lines:
        line = line.strip()
        if line.startswith('Meta Title Recommendation:'):
            section = 'meta_title'
            recommendations['meta_title'] = line.replace('Meta Title Recommendation:', '').strip()
        elif line.startswith('Meta Description Recommendation:'):
            section = 'meta_description'
            recommendations['meta_description'] = line.replace('Meta Description Recommendation:', '').strip()
        elif line.startswith('H1 Tag Recommendation:'):
            section = 'h1_tag'
            recommendations['h1_tag'] = line.replace('H1 Tag Recommendation:', '').strip()
        elif line.startswith('Content Recommendations:'):
            section = 'content'
        elif line.startswith('Recommendation #'):
            if content_rec:
                recommendations['content'].append(content_rec)
            content_rec = {'title': line, 'details': {}}
        elif line.startswith('Rearrangement Suggestion:'):
            if content_rec:
                recommendations['content'].append(content_rec)
            content_rec = {'title': line, 'details': {}}
            section = 'rearrangement'
        elif line.startswith('Provide a final summary'):
            if content_rec:
                if section == 'content':
                    recommendations['content'].append(content_rec)
                elif section == 'rearrangement':
                    recommendations['rearrangements'].append(content_rec)
            section = 'summary'
        elif section == 'content' or section == 'rearrangement':
            if ': ' in line:
                key, value = line.split(': ', 1)
                content_rec['details'][key.strip()] = value.strip()
            else:
                # Handle continuation lines
                if content_rec and 'Content Description' in content_rec['details']:
                    content_rec['details']['Content Description'] += ' ' + line
        elif section == 'summary':
            recommendations['summary'] += line + ' '
    if content_rec and ('title' in content_rec):
        if section == 'content':
            recommendations['content'].append(content_rec)
        elif section == 'rearrangement':
            recommendations['rearrangements'].append(content_rec)
    return recommendations

# Streamlit UI
st.write("Enter your API key and target keyword below:")

openai_api_key = st.text_input("OpenAI API key:", value=st.session_state.openai_api_key, type="password")
keyword = st.text_input("Target keyword:", value=st.session_state.keyword)

# Update session state
st.session_state.openai_api_key = openai_api_key
st.session_state.keyword = keyword

# File uploader for user's HTML file
uploaded_file = st.file_uploader("Upload your HTML file:", type=['html', 'htm'])

# File uploader for competitor HTML files
uploaded_competitor_files = st.file_uploader(
    "Upload competitor HTML files (you can select multiple files):",
    type=['html', 'htm'],
    accept_multiple_files=True
)

if st.button("Optimize Content"):
    if openai_api_key and keyword and uploaded_file and uploaded_competitor_files:
        with st.spinner("Processing..."):
            try:
                # Read user's HTML content
                html_content = uploaded_file.read().decode('utf-8')
                if not html_content:
                    st.error("Uploaded HTML content is empty.")
                else:
                    # Extract content structure and meta info from user's content
                    meta_title, meta_description, h1_text, user_structure = extract_content_structure(html_content)

                    # Analyze competitor content
                    competitor_data = analyze_competitor_content(uploaded_competitor_files)

                    # Prepare data for the prompt
                    user_structure_str = '\n'.join([f"{item['level']}: {item['text']}" for item in user_structure])

                    # Generate detailed recommendations
                    recommendations_text = generate_detailed_recommendations(
                        keyword=keyword,
                        meta_title=meta_title,
                        meta_description=meta_description,
                        h1_text=h1_text,
                        user_structure=user_structure_str,
                        competitor_data=competitor_data,
                        api_key=openai_api_key
                    )

                    if recommendations_text:
                        st.subheader("Detailed Recommendations:")
                        st.text(recommendations_text)

                        # Parse recommendations for Word document
                        recommendations = parse_recommendations(recommendations_text)

                        # Create a Word document with the recommendations
                        doc = Document()
                        doc.add_heading("Content Optimization Recommendations", level=1)
                        doc.add_paragraph(f"Keyword: {keyword}")

                        # Meta Title Recommendation
                        doc.add_heading("Meta Title Recommendation", level=2)
                        doc.add_paragraph(recommendations['meta_title'])

                        # Meta Description Recommendation
                        doc.add_heading("Meta Description Recommendation", level=2)
                        doc.add_paragraph(recommendations['meta_description'])

                        # H1 Tag Recommendation
                        doc.add_heading("H1 Tag Recommendation", level=2)
                        doc.add_paragraph(recommendations['h1_tag'])

                        # Content Recommendations
                        if recommendations['content']:
                            doc.add_heading("Content Recommendations", level=2)
                            for rec in recommendations['content']:
                                doc.add_heading(rec.get('title', ''), level=3)
                                details = rec.get('details', {})
                                for key, value in details.items():
                                    doc.add_heading(f"{key}:", level=4)
                                    doc.add_paragraph(value)

                        # Rearrangement Suggestions
                        if recommendations['rearrangements']:
                            doc.add_heading("Rearrangement Suggestions", level=2)
                            for rec in recommendations['rearrangements']:
                                doc.add_heading(rec.get('title', ''), level=3)
                                details = rec.get('details', {})
                                for key, value in details.items():
                                    doc.add_heading(f"{key}:", level=4)
                                    doc.add_paragraph(value)

                        # Final Summary
                        if recommendations['summary']:
                            doc.add_heading("Final Summary", level=2)
                            doc.add_paragraph(recommendations['summary'])

                        # Create a BytesIO buffer and save the docx content
                        bio = BytesIO()
                        doc.save(bio)
                        bio.seek(0)

                        st.download_button(
                            label="Download Recommendations",
                            data=bio,
                            file_name=f"recommendations_{keyword.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Failed to generate recommendations. Please try again.")
            except Exception as e:
                st.error(f"Error processing files: {str(e)}")
    else:
        st.error("Please enter your API key, target keyword, and upload your HTML files to proceed.")
